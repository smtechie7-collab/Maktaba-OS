import os
import sys
import json

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError as exc:
    DOCX_AVAILABLE = False
    DOCX_ERROR = exc

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '../..')))
from src.core.config import load_config
from src.data.database import DatabaseManager

def enforce_rtl_paragraph(paragraph):
    """Injects <w:bidi/> tags at the paragraph level for strict RTL layout in MS Word."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def enforce_run_font_and_rtl(run, font_name, is_rtl=False):
    """Maps Maktaba styles directly to MS Word XML font and run settings."""
    run.font.name = font_name
    r = run._r
    rPr = r.get_or_add_rPr()
    
    if is_rtl:
        rtl_font = OxmlElement('w:rFonts')
        rtl_font.set(qn('w:cs'), font_name)
        rtl_font.set(qn('w:ascii'), font_name)
        rtl_font.set(qn('w:hAnsi'), font_name)
        rtl_font.set(qn('w:eastAsia'), font_name)
        rtl_font.set(qn('w:hint'), 'cs')
        rPr.append(rtl_font)
        
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        rPr.append(bidi)
    else:
        run.font.name = font_name

class DOCXGenerator:
    def __init__(self, db_path=None):
        config = load_config()
        self.db = DatabaseManager(str(db_path or config.db_path))

    def generate_docx(self, book_id: int, output_path: str, styles: dict = None):
        if not DOCX_AVAILABLE:
            raise RuntimeError("The 'python-docx' library is required for DOCX export. Please run 'pip install python-docx'.")
        
        if not styles:
            styles = {"fonts": {}}

        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT title, author FROM Books WHERE id = ?", (book_id,))
            book_info = cursor.fetchone()
            if not book_info: return

        doc = Document()
        
        # Build Book Metadata (Title Page)
        title_p = doc.add_heading(book_info['title'] or "Untitled Book", 0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if book_info['author']:
            author_p = doc.add_paragraph(f"By {book_info['author']}")
            author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        doc.add_page_break()

        content_blocks = self.db.get_book_content(book_id)
        current_chapter = None

        # Font configuration from Maktaba Styles
        fonts = styles.get("fonts", {})
        ar_font = fonts.get("arabic", "Amiri")
        ar_size = fonts.get("arabic_size", 24)
        ur_font = fonts.get("urdu", "Jameel Noori Nastaliq")
        ur_size = fonts.get("urdu_size", 20)
        en_font = fonts.get("english", "Times New Roman")
        
        for block in content_blocks:
            if block['chapter_title'] != current_chapter:
                current_chapter = block['chapter_title']
                doc.add_heading(current_chapter, level=1)

            data = json.loads(block['content_data'])

            # Arabic Process
            if data.get('ar'):
                p = doc.add_paragraph()
                enforce_rtl_paragraph(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(data['ar'])
                run.font.size = Pt(ar_size)
                enforce_run_font_and_rtl(run, ar_font, is_rtl=True)

            # Urdu Process
            if data.get('ur'):
                p = doc.add_paragraph()
                enforce_rtl_paragraph(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(data['ur'])
                run.font.size = Pt(ur_size)
                enforce_run_font_and_rtl(run, ur_font, is_rtl=True)

            # English Process
            if data.get('en'):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(data['en'])
                run.font.size = Pt(12)
                enforce_run_font_and_rtl(run, en_font, is_rtl=False)

            # Simple separator between layout blocks
            doc.add_paragraph()

        # Output execution
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            
        doc.save(output_path)